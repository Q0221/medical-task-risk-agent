"""图表统计服务：趋势、类型分布、风险分布、负责人排行。"""

from __future__ import annotations

from datetime import date, datetime, timedelta, timezone
from typing import Optional

from sqlalchemy import case, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.enums import ReviewStatus, RiskLevel, TaskStatus, TaskType
from app.models.notification import Notification
from app.models.task import Task
from app.models.user import User
from app.schemas.reports import (
    AssigneeRankItem,
    AssigneeRankResponse,
    ReportDetail,
    ReportHistoryItem,
    ReportHistoryResponse,
    RiskDistItem,
    RiskDistResponse,
    TrendPoint,
    TrendResponse,
    TypeDistItem,
    TypeDistResponse,
)

_TYPE_LABELS: dict[str, str] = {
    TaskType.CUSTOMER_FOLLOWUP: "客户跟进",
    TaskType.PRODUCT_FEEDBACK: "产品反馈",
    TaskType.COMPLAINT: "投诉处理",
    TaskType.ADVERSE_EVENT: "不良事件",
    TaskType.DEVICE_ANOMALY: "设备异常",
    TaskType.COMPLIANCE_REVIEW: "合规审核",
    TaskType.KNOWLEDGE_MAINTAIN: "知识维护",
    "other": "其他",
}

_RISK_LABELS: dict[str, str] = {
    RiskLevel.CRITICAL: "紧急风险",
    RiskLevel.HIGH: "高风险",
    RiskLevel.MEDIUM: "中风险",
    RiskLevel.LOW: "低风险",
}

_KIND_LABELS: dict[str, str] = {
    "daily_summary": "日报",
    "weekly_summary": "周报",
}

# 已审核的 review_status 值集合
_REVIEWED_STATUSES = {
    ReviewStatus.APPROVED,
    ReviewStatus.REJECTED,
    ReviewStatus.ESCALATED,
}


def _utc_midnight(d: date) -> datetime:
    return datetime(d.year, d.month, d.day, tzinfo=timezone.utc)


# ---------------------------------------------------------------------------
# 趋势图
# ---------------------------------------------------------------------------

async def get_trend_data(session: AsyncSession, days: int = 14) -> TrendResponse:
    today = date.today()
    start_dt = _utc_midnight(today - timedelta(days=days - 1))

    rows = (
        await session.execute(
            select(
                func.date(Task.created_at).label("day"),
                func.count(Task.id).label("created"),
                func.sum(
                    case((Task.status == TaskStatus.COMPLETED, 1), else_=0)
                ).label("completed"),
                func.sum(
                    case((Task.status == TaskStatus.OVERDUE, 1), else_=0)
                ).label("overdue"),
            )
            .where(Task.deleted_at.is_(None), Task.created_at >= start_dt)
            .group_by(func.date(Task.created_at))
            .order_by(func.date(Task.created_at))
        )
    ).all()

    data_by_day: dict[str, tuple] = {str(row.day): row for row in rows}

    points: list[TrendPoint] = []
    for offset in range(days):
        target = today - timedelta(days=days - 1 - offset)
        key = str(target)
        row = data_by_day.get(key)
        points.append(
            TrendPoint(
                date=target.strftime("%m-%d"),
                created=int(row.created) if row else 0,
                completed=int(row.completed) if row else 0,
                overdue=int(row.overdue) if row else 0,
            )
        )

    return TrendResponse(points=points, days=days)


# ---------------------------------------------------------------------------
# 任务类型分布
# ---------------------------------------------------------------------------

async def get_type_dist(
    session: AsyncSession,
    date_start: datetime,
    date_end: datetime,
) -> TypeDistResponse:
    rows = (
        await session.execute(
            select(Task.type, func.count(Task.id).label("cnt"))
            .where(
                Task.deleted_at.is_(None),
                Task.created_at >= date_start,
                Task.created_at < date_end,
            )
            .group_by(Task.type)
            .order_by(func.count(Task.id).desc())
        )
    ).all()

    total = sum(row.cnt for row in rows)
    max_count = max((row.cnt for row in rows), default=1)

    items = [
        TypeDistItem(
            type=row.type,
            label=_TYPE_LABELS.get(row.type, row.type),
            count=row.cnt,
            pct=round(row.cnt / max_count * 100, 1),
        )
        for row in rows
    ]
    return TypeDistResponse(items=items, total=total)


# ---------------------------------------------------------------------------
# 风险分布
# ---------------------------------------------------------------------------

async def get_risk_dist(
    session: AsyncSession,
    date_start: datetime,
    date_end: datetime,
) -> RiskDistResponse:
    rows = (
        await session.execute(
            select(
                Task.risk_level,
                func.count(Task.id).label("total"),
                func.sum(
                    case(
                        (
                            Task.review_status.in_(
                                [s.value for s in _REVIEWED_STATUSES]
                            ),
                            1,
                        ),
                        else_=0,
                    )
                ).label("reviewed"),
            )
            .where(
                Task.deleted_at.is_(None),
                Task.risk_level.isnot(None),
                Task.created_at >= date_start,
                Task.created_at < date_end,
            )
            .group_by(Task.risk_level)
        )
    ).all()

    level_order = [
        RiskLevel.CRITICAL,
        RiskLevel.HIGH,
        RiskLevel.MEDIUM,
        RiskLevel.LOW,
    ]

    data_by_level = {row.risk_level: row for row in rows}
    items: list[RiskDistItem] = []
    total = 0
    for level in level_order:
        row = data_by_level.get(level.value)
        count = int(row.total) if row else 0
        reviewed = int(row.reviewed) if row else 0
        total += count
        items.append(
            RiskDistItem(
                level=level.value,
                label=_RISK_LABELS[level],
                count=count,
                reviewed=reviewed,
            )
        )

    return RiskDistResponse(items=items, total=total)


# ---------------------------------------------------------------------------
# 负责人排行
# ---------------------------------------------------------------------------

async def get_assignee_rank(
    session: AsyncSession,
    date_start: datetime,
    date_end: datetime,
    top_n: int = 10,
) -> AssigneeRankResponse:
    rows = (
        await session.execute(
            select(
                User.name,
                func.count(Task.id).label("total"),
                func.sum(
                    case((Task.status == TaskStatus.COMPLETED, 1), else_=0)
                ).label("completed"),
                func.sum(
                    case((Task.status == TaskStatus.OVERDUE, 1), else_=0)
                ).label("overdue"),
            )
            .join(User, Task.assignee_id == User.id)
            .where(
                Task.deleted_at.is_(None),
                Task.created_at >= date_start,
                Task.created_at < date_end,
            )
            .group_by(Task.assignee_id, User.name)
            .order_by(func.count(Task.id).desc())
            .limit(top_n)
        )
    ).all()

    items = [
        AssigneeRankItem(
            name=row.name,
            total=int(row.total),
            completed=int(row.completed),
            overdue=int(row.overdue),
            completion_rate=int(row.completed / row.total * 100) if row.total else 0,
        )
        for row in rows
    ]
    return AssigneeRankResponse(items=items)


# ---------------------------------------------------------------------------
# 历史报告列表
# ---------------------------------------------------------------------------

async def list_report_history(
    session: AsyncSession,
    page: int = 1,
    page_size: int = 20,
    kind: Optional[str] = None,
) -> ReportHistoryResponse:
    base_filters = [
        Notification.deleted_at.is_(None),
        Notification.kind.in_(["daily_summary", "weekly_summary"]),
    ]
    if kind and kind in ("daily_summary", "weekly_summary"):
        base_filters.append(Notification.kind == kind)

    count_result = await session.execute(
        select(func.count(Notification.id)).where(*base_filters)
    )
    total = count_result.scalar_one()

    query = select(Notification).where(*base_filters)

    rows = (
        await session.execute(
            query
            .order_by(Notification.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).scalars().all()

    items = [
        ReportHistoryItem(
            id=notif.id,
            kind=notif.kind,
            kind_label=_KIND_LABELS.get(notif.kind, notif.kind),
            title=notif.title or "",
            preview=(notif.content or "")[:80],
            created_at=notif.created_at,
            status=notif.status,
        )
        for notif in rows
    ]
    return ReportHistoryResponse(items=items, total=total, page=page, page_size=page_size)


async def get_report_detail(session: AsyncSession, report_id: int) -> Optional[ReportDetail]:
    notif = await session.get(Notification, report_id)
    if notif is None or notif.kind not in ("daily_summary", "weekly_summary"):
        return None
    return ReportDetail(
        id=notif.id,
        kind=notif.kind,
        title=notif.title or "",
        content=notif.content or "",
        created_at=notif.created_at,
        status=notif.status,
    )


# ---------------------------------------------------------------------------
# Word / PDF 内容生成
# ---------------------------------------------------------------------------

def build_word_bytes(title: str, content: str) -> bytes:
    """生成 Word .docx 文件字节流。"""
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # 正文段落
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        # 以 ① ② 等标号开头的行视为小标题
        if stripped and stripped[0] in "①②③④⑤⑥⑦⑧⑨⑩":
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)
        else:
            para = doc.add_paragraph(stripped)
            para.style.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_bytes(title: str, content: str) -> bytes:
    """生成 PDF 文件字节流，使用 CID 字体支持中文。"""
    from io import BytesIO

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    base_styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChTitle",
        parent=base_styles["Title"],
        fontName="STSong-Light",
        fontSize=16,
        spaceAfter=14,
        alignment=1,  # center
    )
    section_style = ParagraphStyle(
        "ChSection",
        parent=base_styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        textColor=(0x2E / 255, 0x5E / 255, 0xA8 / 255),
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ChBody",
        parent=base_styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=16,
        spaceAfter=4,
    )

    story = [Paragraph(title, title_style), Spacer(1, 6 * mm)]

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 3 * mm))
            continue
        # 转义 XML 特殊字符
        safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if stripped[0] in "①②③④⑤⑥⑦⑧⑨⑩":
            story.append(Paragraph(f"<b>{safe}</b>", section_style))
        else:
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()
